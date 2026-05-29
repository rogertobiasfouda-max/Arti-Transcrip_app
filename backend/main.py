import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import List

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from openai import OpenAI

# Configurer pydub pour utiliser le binaire ffmpeg embarqué dans imageio-ffmpeg
# (évite d'avoir à installer ffmpeg via apt-get sur le serveur)
try:
    import imageio_ffmpeg
    import pydub
    pydub.AudioSegment.converter = imageio_ffmpeg.get_ffmpeg_exe()
    pydub.AudioSegment.ffmpeg = imageio_ffmpeg.get_ffmpeg_exe()
    pydub.AudioSegment.ffprobe = imageio_ffmpeg.get_ffmpeg_exe()
except Exception:
    pass  # Fallback : pydub cherchera ffmpeg dans le PATH

app = FastAPI(title="Actiwork Transcription API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

MAX_CHUNK_BYTES = 24 * 1024 * 1024   # 24 MB — limite Whisper API
MAX_UPLOAD_BYTES = 200 * 1024 * 1024  # 200 MB — limite upload acceptée

MOIS = [
    "janvier", "février", "mars", "avril", "mai", "juin",
    "juillet", "août", "septembre", "octobre", "novembre", "décembre",
]


# ── Helpers ────────────────────────────────────────────────────────────────

def format_timestamp(seconds: float) -> str:
    """Formater en (m:ss) ou (h:mm:ss)."""
    s = int(seconds)
    if s < 3600:
        return f"({s // 60}:{s % 60:02d})"
    return f"({s // 3600}:{(s % 3600) // 60:02d}:{s % 60:02d})"


def format_srt_ts(seconds: float) -> str:
    """Formater en HH:MM:SS,mmm pour SRT."""
    ms = int(round((seconds % 1) * 1000))
    s = int(seconds)
    return f"{s // 3600:02d}:{(s % 3600) // 60:02d}:{s % 60:02d},{ms:03d}"


def split_audio_if_needed(path: str, file_size: int) -> List[dict]:
    """Découper le fichier audio si > 24 Mo, sinon retourner tel quel."""
    if file_size <= MAX_CHUNK_BYTES:
        return [{"path": path, "offset": 0.0, "temp": False}]

    import pydub

    # S'assurer que pydub utilise le bon ffmpeg (imageio-ffmpeg)
    try:
        import imageio_ffmpeg
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        pydub.AudioSegment.converter = ffmpeg_exe
        pydub.AudioSegment.ffmpeg    = ffmpeg_exe
        print(f"[split] Using ffmpeg: {ffmpeg_exe}")
    except Exception as e:
        print(f"[split] imageio_ffmpeg unavailable: {e}")

    try:
        audio = pydub.AudioSegment.from_file(path)
    except Exception as exc:
        raise HTTPException(
            422,
            f"Impossible de lire/découper ce fichier audio ({len(path)} o). "
            f"Vérifiez que le format est supporté. Détail : {exc}",
        )

    duration_ms = len(audio)
    ratio = MAX_CHUNK_BYTES / file_size
    chunk_ms = int(duration_ms * ratio * 0.9)
    chunks, offset = [], 0

    while offset < duration_ms:
        end = min(offset + chunk_ms, duration_ms)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
        audio[offset:end].export(tmp.name, format="mp3")
        tmp.close()
        chunks.append({"path": tmp.name, "offset": offset / 1000.0, "temp": True})
        offset = end

    return chunks


def transcribe_chunk(path: str, offset: float) -> List[dict]:
    """Transcrire un morceau audio via Whisper, décaler les timestamps."""
    with open(path, "rb") as f:
        resp = client.audio.transcriptions.create(
            model="whisper-1",
            file=f,
            language="fr",
            response_format="verbose_json",
            timestamp_granularities=["segment"],
        )

    return [
        {
            "start": seg.start + offset,
            "end": seg.end + offset,
            "text": seg.text.strip(),
            "timestamp_label": format_timestamp(seg.start + offset),
        }
        for seg in resp.segments
    ]


def build_paragraphs(segments: List[dict], pause: float = 2.0, max_group: int = 5) -> list:
    """Regrouper les segments en paragraphes (pause > 2s ou max 5 segments)."""
    paras, cur = [], []
    for i, seg in enumerate(segments):
        cur.append({"ts": seg["timestamp_label"], "text": seg["text"]})
        last = i == len(segments) - 1
        big_pause = not last and (segments[i + 1]["start"] - seg["end"]) > pause
        if last or big_pause or len(cur) >= max_group:
            paras.append(cur)
            cur = []
    return paras


def build_txt(title: str, date: str, paragraphs: list) -> str:
    lines = [title, date, ""]
    for para in paragraphs:
        lines.append(" ".join(f"{s['ts']} {s['text']}" for s in para))
        lines.append("")
    return "\n".join(lines)


def build_md(title: str, date: str, paragraphs: list) -> str:
    lines = [f"# {title}", f"_{date}_", ""]
    for para in paragraphs:
        lines.append(" ".join(f"**{s['ts']}** {s['text']}" for s in para))
        lines.append("")
    return "\n".join(lines)


def build_srt(segments: List[dict]) -> str:
    blocks = []
    for i, s in enumerate(segments, 1):
        blocks.append(
            f"{i}\n{format_srt_ts(s['start'])} --> {format_srt_ts(s['end'])}\n{s['text']}"
        )
    return "\n\n".join(blocks)


def build_docx_b64(title: str, date: str, paragraphs: list) -> str:
    """Générer un fichier Word (.docx) et retourner en base64."""
    import base64, io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Titre
    h = doc.add_heading(title, level=1)
    h.runs[0].font.color.rgb = RGBColor(0xE2, 0x23, 0x1A)

    # Date
    p_date = doc.add_paragraph(date)
    p_date.runs[0].font.size = Pt(10)
    p_date.runs[0].font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)

    doc.add_paragraph()

    # Paragraphes avec horodatages
    for para in paragraphs:
        p = doc.add_paragraph()
        for seg in para:
            run_ts = p.add_run(seg["ts"] + " ")
            run_ts.font.size = Pt(9)
            run_ts.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)
            run_ts.font.bold = True
            run_text = p.add_run(seg["text"] + " ")
            run_text.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode()


def build_pdf_b64(title: str, date: str, paragraphs: list) -> str:
    """Générer un fichier PDF et retourner en base64."""
    import base64
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Titre
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(226, 35, 26)
    pdf.cell(0, 10, title, ln=True)

    # Date
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(138, 138, 138)
    pdf.cell(0, 6, date, ln=True)
    pdf.ln(4)

    # Contenu
    for para in paragraphs:
        for seg in para:
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(138, 138, 138)
            pdf.write(5, seg["ts"] + " ")
            pdf.set_font("Helvetica", "", 11)
            pdf.set_text_color(26, 26, 26)
            # Encoder en latin-1 pour fpdf (remplacer les caractères non supportés)
            safe = seg["text"].encode("latin-1", errors="replace").decode("latin-1")
            pdf.write(5, safe + " ")
        pdf.ln(6)

    return base64.b64encode(pdf.output()).decode()


# ── Endpoints ──────────────────────────────────────────────────────────────

@app.post("/transcribe")
async def transcribe_audio(file: UploadFile = File(...)):
    allowed = {".mp3", ".wav", ".m4a", ".ogg", ".webm", ".flac", ".mp4"}
    ext = Path(file.filename).suffix.lower()

    if ext not in allowed:
        raise HTTPException(
            400,
            f"Format non supporté. Formats acceptés : {', '.join(sorted(allowed))}",
        )

    content = await file.read()

    # Vérification taille
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            413,
            f"Fichier trop volumineux ({len(content) // (1024*1024)} Mo). "
            f"Limite : {MAX_UPLOAD_BYTES // (1024*1024)} Mo."
        )

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext or ".mp3") as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        chunks = split_audio_if_needed(tmp_path, len(content))
        segments: List[dict] = []

        for chunk in chunks:
            segments.extend(transcribe_chunk(chunk["path"], chunk["offset"]))
            if chunk["temp"]:
                os.unlink(chunk["path"])

        title = Path(file.filename).stem
        dt = datetime.now()
        date = f"{dt.day} {MOIS[dt.month - 1]} {dt.year}, {dt.hour:02d}:{dt.minute:02d}"
        paragraphs = build_paragraphs(segments)

        return {
            "title": title,
            "date": date,
            "segments": segments,
            "paragraphs": paragraphs,
            "exports": {
                "txt":      build_txt(title, date, paragraphs),
                "md":       build_md(title, date, paragraphs),
                "srt":      build_srt(segments),
                "docx_b64": build_docx_b64(title, date, paragraphs),
                "pdf_b64":  build_pdf_b64(title, date, paragraphs),
            },
        }

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(500, str(exc))
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ── Servir le frontend buildé ───────────────────────────────────────────────
_dist = Path(__file__).parent / "frontend" / "dist"
if _dist.exists():
    app.mount("/", StaticFiles(directory=str(_dist), html=True), name="static")


# ── Point d'entrée (Railway lit PORT directement via os.environ) ────────────
if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8080))
    uvicorn.run(app, host="0.0.0.0", port=port)
